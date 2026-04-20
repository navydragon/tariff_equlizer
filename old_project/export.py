import os
from docx import Document

def collect_python_files_to_docx(directory, output_file):
    # Создаем новый docx-документ
    document = Document()

    # Проходим по всем файлам и подпапкам в указанной директории
    for root, dirs, files in os.walk(directory):
        if '.venv' in dirs:
            dirs.remove('.venv')
        for file in files:
            if file.endswith(".py"):
                file_path = os.path.join(root, file)
                print(f"Обрабатывается файл: {file_path}")
                # Добавляем название файла в документ
                document.add_heading(f"Файл: {file}", level=2)

                # Читаем содержимое файла
                try:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        content = f.read()

                    # Добавляем содержимое файла в документ
                    document.add_paragraph(content)
                except Exception as e:
                    paragraph = document.add_paragraph(f"Не удалось прочитать файл: {e}")
                    paragraph.runs[0].italic = True

                    # Добавляем разделитель между файлами
                document.add_paragraph("\n---\n")

    # Сохраняем документ
    document.save(output_file)

# Укажите путь к папке и имя выходного файла
folder_path = "C:/Users/User/PycharmProjects/tariff_equlizer/"
output_docx = "код.docx"

collect_python_files_to_docx(folder_path, output_docx)

print(f"Содержимое всех файлов .py собрано в {output_docx}")
